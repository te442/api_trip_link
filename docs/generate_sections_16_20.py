# -*- coding: utf-8 -*-
"""Generates sections 16–20 Word document — UI Design Chapter for TripLink."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\User\Desktop\פרויקט טיול\API_trip_link\docs\16_20_תכנון_ממשק_משתמש.docx"


# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            set_rtl(p)
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                set_rtl(p)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_ascii_block(doc, lines, font_size=9):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(font_size)


def add_screen_section(doc, num, title, route, component, role_bullets, structure_bullets, components_rows):
    add_heading(doc, f"17.{num} {title}", level=3)
    add_para(doc, f"נתיב: {route} | Component: {component}")

    add_heading(doc, f"17.{num}.1 תפקיד המסך", level=4)
    for b in role_bullets:
        add_bullet(doc, b)

    add_heading(doc, f"17.{num}.2 תיאור מבנה המסך", level=4)
    add_para(doc, "התיאור הבא מתייחס למבנה לוגי ולסידור רכיבים — ללא עיצוב גרפי (צבעים, לוגו).")
    for b in structure_bullets:
        add_bullet(doc, b)

    add_heading(doc, f"17.{num}.3 פירוט רכיבי התצוגה", level=4)
    add_table(doc,
        ["שם הרכיב", "סוג", "תפקיד", "פעולת משתמש", "השפעה על המערכת"],
        components_rows)


# ─── Section 16 ─────────────────────────────────────────────────────────────

def build_section_16(doc):
    add_heading(doc, "16. Screen Flow Diagram — תרשים מסכים והמעברים ביניהם", level=1)

    add_heading(doc, "16.1 מבוא", level=2)
    add_para(doc, "פרק זה מתאר את מבנה המסכים באפליקציית הלקוח (Frontend) של מערכת TripLink, את ההיררכיה ביניהם ואת המעברים (Navigation Flow) במהלך שימוש במערכת. האפליקציה בנויה כ-Single Page Application (SPA) ב-Angular 19, עם ניווט מבוסס Router, אימות JWT ו-AuthGuard על מסכים פנימיים.")
    add_para(doc, "מערכת TripLink תומכת בתהליך מלא: התחברות, תכנון טיול (אשף רב-שלבי), חישוב מסלול אופטימלי והצגת תוצאות הכוללת מפה, פירוט תחבורה ויעדים.")

    add_heading(doc, "16.2 מבנה המסכים והיררכיה", level=2)
    add_para(doc, "האפליקציה כוללת שני מסכי אימות, מסך מעטפת (Shell) וחמישה מסכי תוכן. מסכים פנימיים מוגנים על ידי AuthGuard.")

    add_table(doc,
        ["#", "שם המסך", "נתיב (Route)", "Component", "תפקיד"],
        [
            ["—", "מסך מעטפת", "—", "AppComponent", "תפריט ניווט + router-outlet"],
            ["0א", "התחברות", "/login", "LoginComponent", "כניסה עם JWT"],
            ["0ב", "הרשמה", "/register", "RegisterComponent", "יצירת משתמש"],
            ["1", "הטיולים שלי", "/my-trips", "MyTripsComponent", "ניהול טיולים"],
            ["2", "אשף תכנון", "/plan", "TripWizardComponent", "תכנון ב-3 שלבים"],
            ["3", "חישוב מסלול", "/plan/optimize/:tripId", "OptimizeScreenComponent", "הרצת אופטימיזציה"],
            ["4", "תוצאות הטיול", "/trips/:id/result", "TripResultComponent", "הצגת מסלול מלא"],
            ["5", "רשימת יעדים", "/destinations", "DestinationsListComponent", "צפייה ביעדים"],
        ])

    add_heading(doc, "16.2.1 היררכיית המסכים", level=3)
    add_ascii_block(doc, [
        "TripLink (Angular SPA)",
        "├── אזור ציבורי: /login, /register",
        "├── מעטפת (App Shell): NavBar + router-outlet",
        "├── /my-trips — הטיולים שלי",
        "├── /plan — אשף תכנון (שלבים 1–3)",
        "├── /plan/optimize/:tripId — חישוב מסלול",
        "├── /trips/:id/result — תוצאות (מפה + timeline)",
        "└── /destinations — רשימת יעדים",
    ])

    add_heading(doc, "16.2.2 תרשים מעברים בין מסכים", level=3)
    add_ascii_block(doc, [
        "       כניסה (/) → redirect → /login",
        "       /login ──הצלחה──► /my-trips ◄──הצלחה── /register",
        "       /my-trips ──תכנן טיול──► /plan (3 שלבים)",
        "       /plan ──צור טיול──► /plan/optimize/:id",
        "       /optimize ──חשב מסלול──► /trips/:id/result",
        "       NavBar: מעבר חופשי בין /my-trips, /plan, /destinations",
        "       AuthGuard: גישה ללא JWT → /login",
    ], font_size=10)

    add_heading(doc, "16.3 זרימת משתמש עיקרית", level=2)
    for step in [
        "כניסה לאפליקציה → מסך התחברות.",
        "הרשמה או התחברות → JWT נשמר → מסך הטיולים שלי.",
        "תכנון טיול חדש → אשף 3 שלבים (פרטים, העדפות, זמנים).",
        "יצירת טיול ב-API → מסך חישוב מסלול.",
        "הרצת אופטימיזציה → מסך תוצאות עם מפה ופירוט.",
        "חזרה לרשימת טיולים או תכנון טיול נוסף.",
    ]:
        add_bullet(doc, step)

    add_heading(doc, "16.4 טבלת מעברים", level=2)
    add_table(doc,
        ["מסך מקור", "פעולה", "מסך יעד", "סוג מעבר"],
        [
            ["/", "Redirect", "/login", "Router"],
            ["/login", "התחברות מוצלחת", "/my-trips", "navigate"],
            ["/login", "הירשם כאן", "/register", "RouterLink"],
            ["/register", "הרשמה מוצלחת", "/my-trips", "navigate"],
            ["/my-trips", "+ תכנן טיול", "/plan", "RouterLink"],
            ["/plan", "צור טיול", "/plan/optimize/:id", "navigate"],
            ["/optimize", "חשב מסלול", "/trips/:id/result", "navigate"],
            ["/result", "חזור", "/my-trips", "RouterLink"],
            ["NavBar", "התנתק", "/login", "logout"],
            ["AuthGuard", "ללא JWT", "/login", "Guard"],
        ])


# ─── Section 17 ─────────────────────────────────────────────────────────────

def build_section_17(doc):
    add_heading(doc, "17. תיאור מסכי המערכת", level=1)
    add_para(doc, "סעיף זה מפרט את כל מסכי המערכת: תפקיד, מבנה לוגי ורכיבי תצוגה. התיאור מתמקד בפונקציונליות ולא בעיצוב גרפי.")

    # 17.0 Shell
    add_screen_section(doc, "0", "מסך מעטפת (App Shell)", "—", "AppComponent",
        [
            "מטרה: לספק ניווט קבוע בין מסכי המערכת לאחר התחברות.",
            "פעולות: מעבר בין מסכים, התנתקות.",
            "נתונים: מוצג רק כאשר המשתמש מחובר (JWT קיים).",
        ],
        [
            "אזור עליון (Header): תפריט ניווט אופקי.",
            "אזור תוכן (Body): router-outlet — מציג את המסך הפעיל.",
            "כיוון RTL בכל רוחב המסך.",
        ],
        [
            ["NavBar", "Navigation Bar", "ניווט בין מסכים", "לחיצה על קישור", "מעבר ל-Route המתאים"],
            ["router-outlet", "Container", "הצגת מסך פעיל", "—", "טעינת Component לפי Route"],
            ["כפתור התנתק", "Button", "יציאה מהמערכת", "לחיצה", "מחיקת JWT, מעבר ל-/login"],
        ])

    # 17.1 Login
    add_screen_section(doc, "1", "מסך התחברות", "/login", "LoginComponent",
        [
            "מטרה: אימות משתמש קיים וכניסה למערכת.",
            "פעולות: הזנת אימייל וסיסמה, התחברות, מעבר להרשמה.",
            "נתונים נשלחים: POST /api/auth/login.",
            "נתונים מוחזרים: JWT, userId, fullName, email.",
        ],
        [
            "מבנה מרכזי: כרטיס (Card) במרכז המסך.",
            "אזור כותרת: כותרת «התחברות».",
            "אזור טופס: שדות אימייל וסיסמה.",
            "אזור פעולות: כפתור התחברות, קישור להרשמה.",
            "אזור הודעות: הודעת שגיאה (במידת הצורך).",
        ],
        [
            ["שדה אימייל", "TextBox (email)", "הזנת אימייל", "הקלדה", "ערך לשליחה ב-login"],
            ["שדה סיסמה", "TextBox (password)", "הזנת סיסמה", "הקלדה", "ערך לשליחה ב-login"],
            ["כפתור התחבר", "Button", "שליחת טופס", "לחיצה / Enter", "קריאת API, שמירת JWT, מעבר ל-/my-trips"],
            ["קישור הרשמה", "Link", "מעבר להרשמה", "לחיצה", "ניווט ל-/register"],
            ["הודעת שגיאה", "Label", "משוב על כשל", "—", "הצגת טקסט שגיאה מה-API"],
        ])

    # 17.2 Register
    add_screen_section(doc, "2", "מסך הרשמה", "/register", "RegisterComponent",
        [
            "מטרה: יצירת חשבון משתמש חדש.",
            "פעולות: מילוי פרטים, הרשמה, מעבר להתחברות.",
            "נתונים נשלחים: POST /api/auth/register (fullName, email, phone, password).",
        ],
        [
            "מבנה זהה למסך התחברות: Card מרכזי.",
            "שדות: שם מלא, אימייל, טלפון, סיסמה.",
            "כפתור הרשמה וקישור להתחברות.",
        ],
        [
            ["שדה שם מלא", "TextBox", "הזנת שם", "הקלדה", "שמירה ב-Users"],
            ["שדה אימייל", "TextBox (email)", "הזנת אימייל", "הקלדה", "מזהה התחברות"],
            ["שדה טלפון", "TextBox (tel)", "הזנת טלפון", "הקלדה", "פרטי משתמש"],
            ["שדה סיסמה", "TextBox (password)", "הגדרת סיסמה", "הקלדה", "גיבוב ושמירה ב-API"],
            ["כפתור הירשם", "Button", "יצירת חשבון", "לחיצה", "POST register, JWT, מעבר ל-/my-trips"],
            ["קישור התחברות", "Link", "מעבר להתחברות", "לחיצה", "ניווט ל-/login"],
        ])

    # 17.3 My Trips
    add_screen_section(doc, "3", "מסך הטיולים שלי", "/my-trips", "MyTripsComponent",
        [
            "מטרה: הצגת טיולים שנוצרו על ידי המשתמש המחובר.",
            "פעולות: צפייה במסלול, מחיקת טיול, יצירת טיול חדש.",
            "נתונים: GET /api/trips/user/{userId}.",
        ],
        [
            "אזור כותרת: שם המסך + ברכת שלום (שם המשתמש).",
            "אזור פעולות: כפתור «+ תכנן טיול חדש».",
            "אזור תוכן: טבלת טיולים או הודעת «אין טיולים».",
            "מצב טעינה: הודעת «טוען...».",
        ],
        [
            ["ברכת שלום", "Label", "זיהוי משתמש", "—", "הצגת fullName מ-JWT"],
            ["כפתור תכנן טיול", "Link/Button", "יצירת טיול", "לחיצה", "מעבר ל-/plan"],
            ["טבלת טיולים", "Table", "רשימת טיולים", "—", "הצגת tripId, שם, תאריך, כתובת"],
            ["קישור צפה במסלול", "Link", "צפייה בתוצאות", "לחיצה", "מעבר ל-/trips/{id}/result"],
            ["כפתור מחק", "Button", "מחיקת טיול", "לחיצה + confirm", "DELETE /api/trips/{id}"],
            ["הודעת ריק", "Label", "אין נתונים", "—", "«אין טיולים עדיין»"],
        ])

    # 17.4 Wizard
    add_screen_section(doc, "4", "אשף תכנון טיול", "/plan", "TripWizardComponent",
        [
            "מטרה: איסוף כל פרמטרי הטיול לפני יצירתו ב-API.",
            "פעולות: מילוי 3 שלבים, מעבר בין שלבים, יצירת טיול, ביטול.",
            "נתונים נטענים: Lookups (קטגוריות, קושי, מאפיינים, אזורים).",
            "נתונים נשלחים: POST /api/trips (CreateTripDto).",
        ],
        [
            "אזור כותרת: «תכנון טיול חדש».",
            "מחוון שלבים (Step Indicator): 1. פרטים | 2. העדפות | 3. זמנים.",
            "אזור טופס: משתנה לפי שלב נוכחי.",
            "אזור פעולות: הקודם, הבא, צור טיול, ביטול.",
        ],
        [
            ["מחוון שלבים", "Badge/Label", "מיקום באשף", "—", "הדגשת שלב פעיל"],
            ["שם הטיול", "TextBox", "שם הטיול", "הקלדה", "שדה חובה — tripName"],
            ["בחירת אזור", "Dropdown", "אזור גיאוגרafi", "בחירה", "סינון יעדים — region"],
            ["תאריך", "DatePicker", "תאריך טיול", "בחירה", "tripDate"],
            ["כתובת התחלה", "TextBox", "נקודת יציאה", "הקלדה", "addressStart"],
            ["רמת קושי", "Dropdown", "העדפת קושי", "בחירה", "levelId"],
            ["קטגוריות", "Checkbox Grid", "בחירה מרובה", "סימון", "categoryIds"],
            ["מאפיינים", "Checkbox Grid", "דרישות נגישות", "סימון", "featureIds"],
            ["מינ/מקס יעדים", "NumberInput", "כמות יעדים", "הקלדה", "minNumDes, maxNumDes"],
            ["שעות יציאה/חזרה", "TimeInput", "חלון זמן", "בחירה", "startTime, endTime"],
            ["פרמטרי אופטימיזציה", "NumberInput", "הגבלות זמן", "הקלדה", "queryParams לאופטימיזציה"],
            ["כפתור הבא/הקודם", "Button", "ניווט בשלבים", "לחיצה", "שינוי step"],
            ["כפתור צור טיול", "Button", "שמירה", "לחיצה", "POST trip → /plan/optimize/:id"],
        ])

    # 17.5 Optimize
    add_screen_section(doc, "5", "מסך חישוב מסלול", "/plan/optimize/:tripId", "OptimizeScreenComponent",
        [
            "מטרה: הרצת אלגוריתם האופטימיזציה על הטיול שנוצר.",
            "פעולות: צפייה בסיכום, הפעלת חישוב מסלול.",
            "נתונים: POST /api/trips/optimize, POST /api/trips/{id}/save-route.",
        ],
        [
            "אזור כותרת: «חישוב מסלול אופטימלי».",
            "כרטיס סיכום (read-only): שם טיול, כתובת, שעות.",
            "אזור פעולה: כפתור חישוב.",
            "קישור חזרה לתכנון.",
        ],
        [
            ["כרטיס סיכום", "Card", "תצוגת פרטי טיול", "—", "GET /api/trips/{id}"],
            ["כפתור חשב מסלול", "Button", "הפעלת אופטימיזציה", "לחיצה", "POST optimize, שמירה, מעבר ל-result"],
            ["הודעת שגיאה", "Label", "כשל בחישוב", "—", "הצגת error מה-API"],
            ["קישור חזור", "Link", "חזרה לאשף", "לחיצה", "ניווט ל-/plan"],
        ])

    # 17.6 Result
    add_screen_section(doc, "6", "מסך תוצאות הטיול", "/trips/:id/result", "TripResultComponent",
        [
            "מטרה: הצגת המסלול האופטימלי המלא — נקודת הסיום של תהליך התכנון.",
            "פעולות: צפייה במפה, ב-timeline, בפירוט תחבורה; חזרה לטיולים או תכנון חדש.",
            "נתונים: TripStateService (מיידי) או GET /api/trips/{id}/itinerary.",
        ],
        [
            "אזור כותרת: שם טיול, כתובת התחלה, סטטיסטיקות (יעדים, ציון, יעילות).",
            "אזור ראשי — פריסה דו-עמודתית (Grid): מפה | timeline.",
            "במסכים צרים: עמודה אחת (Responsive).",
            "אזור תחתון: סיכום מילולי (Narrative) וכפתורי פעולה.",
        ],
        [
            ["סרגל סטטיסטיקות", "StatsBar", "סיכום מספרי", "—", "destinationCount, totalScore, transitEfficiency"],
            ["מפת Google Maps", "Map", "תצוגה גיאוגרafית", "—", "markers + polyline"],
            ["MapMarker", "Marker", "נקודת יעד", "—", "מיקום לפי lat/lon"],
            ["MapPolyline", "Polyline", "קו מסלול", "—", "חיבור נקודות לפי סדר"],
            ["כרטיס יעד (leg-card)", "Card", "פירוט יעד במסלול", "—", "שם, זמנים, תמונה, תחבורה"],
            ["תמונת יעד", "Image", "תצוגה ויזואלית", "—", "imageUrl מה-API"],
            ["פאנל תחבורה", "Panel", "פירוט אוטובוסים", "—", "קווים, תחנות, זמנים"],
            ["סיכום מילולי", "TextBlock", "Narrative", "—", "טקסט מפורט מה-API"],
            ["כפתורי פעולה", "Link", "ניווט", "לחיצה", "חזרה ל-/my-trips או /plan"],
        ])

    # 17.7 Destinations
    add_screen_section(doc, "7", "מסך רשימת יעדים", "/destinations", "DestinationsListComponent",
        [
            "מטרה: הצגת כל היעדים הזמינים במערכת — מסך עזר לצפייה.",
            "פעולות: צפייה בלבד — אין עריכה.",
            "נתונים: GET /api/destinations.",
        ],
        [
            "אזור כותרת: «יעדים».",
            "אזור תוכן: טבלה או הודעת «אין יעדים».",
        ],
        [
            ["טבלת יעדים", "Table", "רשימת יעדים", "—", "desId, שם, אזור, קושי, מטייל, זמן"],
            ["הודעת טעינה", "Label", "מצב המתנה", "—", "«טוען...»"],
            ["הודעת ריק", "Label", "אין נתונים", "—", "«אין יעדים»"],
        ])


# ─── Section 18 ─────────────────────────────────────────────────────────────

def build_section_18(doc):
    add_heading(doc, "18. הודעות למשתמש", level=1)
    add_para(doc, "מערכת TripLink משתמשת בהודעות inline (בתוך המסך) ובחלון אישור דפדפן (confirm). אין שימוש ב-Toast notifications.")

    add_heading(doc, "18.1 הודעות שגיאה", level=2)
    add_table(doc,
        ["הודעה", "מסך", "מתי מוצגת", "מטרה"],
        [
            ["שגיאה בהתחברות", "התחברות", "כשל POST /api/auth/login", "להודיע על אימייל/סיסמה שגויים"],
            ["שגיאה בהרשמה", "הרשמה", "כשל POST /api/auth/register", "להודיע על אימייל קיים או נתונים חסרים"],
            ["שגיאה ביצירת הטיול", "אשף תכנון", "כשל POST /api/trips", "להודיע על בעיה בשמירה"],
            ["שגיאה בחישוב המסלול", "חישוב מסלול", "כשל POST /api/trips/optimize", "להודיע על כשל באלגוריתם"],
            ["לא נמצא טיול", "חישוב מסלול", "GET trip נכשל", "להודיע שמזהה טיול לא תקין"],
            ["לא נמצא מסלול לטיול זה", "תוצאות", "GET itinerary נכשל", "להנחות להריץ אופטימיזציה קודם"],
        ])

    add_heading(doc, "18.2 הודעות מידע וטעינה", level=2)
    add_table(doc,
        ["הודעה", "מסך", "מתי מוצגת", "מטרה"],
        [
            ["טוען...", "הטיולים שלי / יעדים", "בעת טעינת נתונים", "לסמן המתנה"],
            ["טוען נתונים...", "אשף תכנון", "טעינת Lookups", "לסמן המתנה לפני הצגת טופס"],
            ["טוען מסלול...", "תוצאות", "טעינת itinerary", "לסמן המתנה"],
            ["מתחבר... / נרשם... / יוצר טיול...", "שונים", "בעת שליחת טופס", "לסמן עיבוד"],
            ["מחשב מסלול...", "חישוב מסלול", "בעת אופטימיזציה", "לסמן עיבוד"],
            ["אין טיולים עדיין", "הטיולים שלי", "רשימה ריקה", "להנחות ליצור טיול"],
            ["אין יעדים", "יעדים", "רשימה ריקה", "לציין שאין נתונים"],
            ["הגדר googleMapsApiKey...", "תוצאות", "מפתח מפה חסר", "להנחות להפעלת מפה"],
        ])

    add_heading(doc, "18.3 הודעות אישור (Confirmation)", level=2)
    add_table(doc,
        ["הודעה", "מסך", "מתי מוצגת", "מטרה"],
        [
            ["האם למחוק את הטיול?", "הטיולים שלי", "לחיצה על «מחק»", "לאישור פעולה בלתי הפיכה"],
        ])

    add_heading(doc, "18.4 משוב הצלחה", level=2)
    add_para(doc, "המערכת אינה מציגה הודעות הצלחה מפורשות (Toast). הצלחה מסומנת באמצעות מעבר אוטומטי למסך הבא:")
    add_bullet(doc, "התחברות/הרשמה מוצלחת → מעבר ל-/my-trips.")
    add_bullet(doc, "יצירת טיול מוצלחת → מעבר ל-/plan/optimize/:tripId.")
    add_bullet(doc, "חישוב מסלול מוצלח → מעבר ל-/trips/:id/result.")

    add_heading(doc, "18.5 מדיניות הצגת הודעות", level=2)
    add_bullet(doc, "שגיאות מוצגות כ-Label בתוך המסך, מתחת לטופס או לכפתור.")
    add_bullet(doc, "כפתורי שליחה מושבתים (disabled) בזמן טעינה למניעת שליחה כפולה.")
    add_bullet(doc, "טקסט הכפתור משתנה לזמן הטעינה (למשל: «יוצר טיול...»).")


# ─── Section 19 ─────────────────────────────────────────────────────────────

def build_section_19(doc):
    add_heading(doc, "19. ממשק משתמש (User Interface)", level=1)
    add_para(doc, "סעיף זה מתאר עקרונות פונקציונליים של הממשק. העיצוב הגרפי הסופי (צבעים, לוגו, מיתוג) אינו חלק מגרסת ה-MVP הנוכחית ויוגדר בשלב עיצוב עתידי.")

    add_heading(doc, "19.1 נוחות שימוש (Usability)", level=2)
    add_bullet(doc, "אשף רב-שלבי (Wizard) מפשט את תכנון הטיול — חלוקה ל-3 שלבים לוגיים: פרטים, העדפות, זמנים.")
    add_bullet(doc, "מחוון שלבים (Step Indicator) מציג למשתמש את מיקומו בתהליך.")
    add_bullet(doc, "Placeholders בשדות טקסט מנחים את המשתמש (למשל: «טיול דרום קלאסי»).")
    add_bullet(doc, "כפתורים מושבתים בזמן עיבוד למניעת פעולות כפולות.")
    add_bullet(doc, "שדות חובה מסומנים בכוכבית (*) ונבדקים לפני מעבר לשלב הבא.")

    add_heading(doc, "19.2 אחידות מבנית", level=2)
    add_bullet(doc, "תבנית חוזרת בכל מסך: כותרת → תוכן → אזור פעולות.")
    add_bullet(doc, "מסכי אימות (התחברות, הרשמה) במבנה Card מרכזי זהה.")
    add_bullet(doc, "טבלאות ברשימות (טיולים, יעדים) עם מבנה אחיד: כותרות עמודות + שורות נתונים.")
    add_bullet(doc, "כפתורי פעולה ראשיים ומשניים בעלי תפקיד קבוע בכל מסך.")

    add_heading(doc, "19.3 ניווט במערכת", level=2)
    add_bullet(doc, "NavBar קבוע בראש המסך (לאחר התחברות) מאפשר מעבר חופשי בין מסכים עיקריים.")
    add_bullet(doc, "AuthGuard מפנה משתמש לא מחובר אוטומטית ל-/login.")
    add_bullet(doc, "שלבי האשף משמשים כ-breadcrumb מרומז — המשתמש יודע באיזה שלב בתהליך הוא נמצא.")
    add_bullet(doc, "זרימה ליניארית מוגדרת: תכנון → אופטימיזציה → תוצאות, עם אפשרות חזרה בכל שלב.")

    add_heading(doc, "19.4 כיוון RTL ותמיכה בעברית", level=2)
    add_bullet(doc, "כל המסכים מוגדרים ב-direction: rtl לתמיכה בעברית.")
    add_bullet(doc, "טבלאות וטפסים מיושרים לימין.")
    add_bullet(doc, "תוויות (Labels) בעברית לכל שדות הקלט.")

    add_heading(doc, "19.5 Responsive Design", level=2)
    add_bullet(doc, "מסך תוצאות הטיול משתמש ב-Grid דו-עמודתי (מפה | timeline).")
    add_bullet(doc, "במסכים ברוחב מתחת ל-900px הפריסה עוברת לעמודה אחת.")
    add_bullet(doc, "מסכי אימות ואשף מותאמים לרוחב מסך מלא עם max-width קבוע.")

    add_heading(doc, "19.6 נגישות בסיסית", level=2)
    add_bullet(doc, "לכל שדה קלט יש Label מלווה.")
    add_bullet(doc, "כפתורים מושבתים (disabled) בזמן טעינה.")
    add_bullet(doc, "הגבלה: אין יישום מלא של תקני ARIA — נגישות בסיסית בלבד.")

    add_heading(doc, "19.7 עיצוב גרפי — הערת היקף", level=2)
    add_para(doc, "הגרסה הנוכחית של TripLink מספקת שלד פונקציונלי (Functional Prototype). העיצוב הוויזואלי הסופי — לרבות פלטת צבעים, לוגו, אייקונים וטיפוגרפיה — יוגדר בשלב עיצוב עתידי על ידי המפתח. פרק זה מתמקד במבנה, בזרימה ובפונקציונליות של הממשק.")


# ─── Section 20 ─────────────────────────────────────────────────────────────

def build_section_20(doc):
    add_heading(doc, "20. סיכום ממשק המשתמש", level=1)
    add_para(doc, "ממשק המשתמש של מערכת TripLink תוכנן לתמוך בתהליך תכנון טיול מקצה לקצה, בצורה פשוטה, ברורה ויעילה.")
    add_para(doc, "מבנה המסכים משקף את שלבי העבודה הלוגיים: אימות המשתמש, ניהול טיולים, תכנון (אשף), חישוב מסלול והצגת תוצאות. הפרדה זו מפחיתה את עומס הקוגניטיבי על המשתמש ומאפשרת התמקדות בכל שלב בנפרד.")
    add_para(doc, "אשף התכנון הרב-שלבי מאגד את כל פרמטרי הטיול — פרטים, העדפות וזמנים — בממשק אחד מובנה, תוך טעינת נתוני עזר (Lookups) מה-API. מסך חישוב המסלול מפריד בין הגדרת הטיול לבין הרצת האלגוריתם, ומאפשר למשתמש לוודא את הפרטים לפני החישוב.")
    add_para(doc, "מסך תוצאות הטיול מהווה את נקודת הסיום והמרכזית: הוא משלב תצוגה גיאוגרפית (מפת Google Maps), timeline מפורט של כל יעד, פירוט תחבורה ציבורית וסיכום מילולי. שילוב זה מאפשר למשתמש להבין את המסלול המומלץ הן ברמה חזותית והן ברמת פרטים.")
    add_para(doc, "מערכת ההודעות מספקת משוב ברור על שגיאות, מצבי טעינה ורשימות ריקות. מנגנון האימות (JWT) והגנת המסכים (AuthGuard) מבטיחים שכל משתמש רואה רק את הטיולים שלו.")
    add_para(doc, "לסיכום, ממשק TripLink מספק שלד פונקציונלי מלא לתכנון טיולים, המבוסס על עקרונות נוחות שימוש, אחידות מבנית וניווט ברור. העיצוב הגרפי יושלם בשלב מאוחר, אך המבנה והזרימה שתוארו בפרק זה מהווים את הבסיס לחוויית משתמש יעילה.")


# ─── Build ────────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_heading("ספר פרויקט גמר — TripLink", level=0)
    set_rtl(title)
    add_heading(doc, "פרק: תכנון ממשק משתמש (UI Design)", level=1)
    add_para(doc, "מקצוע: מדעי המחשב | רכיב: אפליקציית Angular — API_trip_link/trip-planner-app")
    add_para(doc, "טכנולוגיות: Angular 19, JWT, Google Maps API, REST API (.NET)")
    doc.add_paragraph()

    build_section_16(doc)
    doc.add_page_break()
    build_section_17(doc)
    doc.add_page_break()
    build_section_18(doc)
    build_section_19(doc)
    build_section_20(doc)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_document()
