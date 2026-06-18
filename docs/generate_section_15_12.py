# -*- coding: utf-8 -*-
"""Generates section 15.12 Word document for TripLink project book."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\User\Desktop\פרויקט טיול\API_trip_link\docs\15.12_תיאור_ארכיטקטוני.docx"


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            set_rtl(p)
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                set_rtl(p)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_module_tree(doc):
    add_heading(doc, "עץ מודולים — מערכת TripLink", level=2)
    add_para(doc, "העץ הבא מציג את מבנה המודולים והיחידות במערכת, מרמת הלקוח ועד בסיס הנתונים ושירותים חיצוניים.")

    tree = [
        "TripLink — מערכת תכנון ואופטימיזציה של טיולים",
        "│",
        "├── שכבת לקוח (Client Layer)",
        "│   └── אפליקציית Angular",
        "│       ├── ממשק משתמש (UI)",
        "│       ├── שליחת בקשות HTTP/JSON",
        "│       └── תצוגת תוצאות אופטימיזציה",
        "│",
        "├── שכבת API (Presentation Layer)",
        "│   ├── TripsController",
        "│   │   ├── GET/POST/DELETE — ניהול טיולים",
        "│   │   ├── POST optimize — הרצת אופטימיזציה",
        "│   │   └── POST save-route — שמירת מסלול",
        "│   ├── UsersController — CRUD משתמשים",
        "│   ├── DestinationsController — שליפת יעדים ותחנות",
        "│   └── LookupsController — נתוני עזר לטופס",
        "│",
        "├── שכבת לוגיקה עסקית (Business Logic Layer)",
        "│   ├── TripService — CRUD טיולים, שמירת מסלול",
        "│   ├── UserService — CRUD משתמשים",
        "│   ├── DestinationService — שליפת יעדים",
        "│   ├── LookupService — מתווך ל-Lookups",
        "│   └── מנוע אופטימיזציה (Optimizer)",
        "│       ├── OptimizerServiceImpl — נקודת כניסה",
        "│       ├── OptimizerPipeline — ניהול שלבים 0–6",
        "│       ├── OptimizeResultMapper — המרה ל-DTO",
        "│       ├── שלבי Pipeline (IOptimizerStep)",
        "│       │   ├── Step0_InputLoader — טעינת נתונים",
        "│       │   ├── Step1_WeightCalculator — חישוב משקלים",
        "│       │   ├── Step2_ScoreTableBuilder — טבלת עלויות",
        "│       │   ├── Step3_DestinationRanker — דירוג יעדים",
        "│       │   ├── Step4_InitialRouteBuilder — מסלול התחלתי",
        "│       │   ├── Step5_SaOptimizer — Simulated Annealing",
        "│       │   └── Step6_TripItineraryBuilder — בניית itinerary",
        "│       └── רכיבי עזר אלגוריתמיים",
        "│           ├── WeightCalculator",
        "│           ├── ArcCostCalculator",
        "│           ├── RouteBuilder",
        "│           └── ScoreTable",
        "│",
        "├── שירות תחבורה (Transit)",
        "│   ├── ITransitApiService — ממשק",
        "│   ├── GoogleMapsTransitApiService — Google Maps API",
        "│   └── MockTransitApiService — גיבוי (Fallback)",
        "│",
        "├── שכבת גישה לנתונים (Data Access Layer)",
        "│   ├── TripContext — DbContext (Entity Framework Core)",
        "│   ├── ILookupRepository / LookupRepository",
        "│   └── IOptimizerDataRepository / OptimizerDataRepository",
        "│",
        "├── שכבת מודל (Models Layer)",
        "│   ├── Entities (17 טבלאות) — Destination, Trip, Station, User...",
        "│   ├── DTOs (14) — TripDto, OptimizeResultDto, CategoryDto...",
        "│   ├── מודלי Optimizer — OptimizationContext, OptimizerRoute...",
        "│   └── GoogleMapsResponseModels — מודלי API חיצוני",
        "│",
        "├── בסיס נתונים (Persistence)",
        "│   └── SQL Server — מסד TripLink",
        "│       ├── SeedData.sql / ClearData.sql",
        "│       └── 16+ טבלאות (Destination, Trip, Station, Bus...)",
        "│",
        "└── שירותים חיצוניים (External Services)",
        "    └── Google Maps Directions API — זמני תחבורה ציבורית",
    ]

    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run("\n".join(tree))
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(12)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("ספר פרויקט גמר — TripLink", level=0)
    set_rtl(title)
    subtitle = doc.add_heading("סעיף 15.12 — תיאור ארכיטקטוני של המערכת", level=1)
    set_rtl(subtitle)
    add_para(doc, "מקצוע: מדעי המחשב | מערכת: TripLink — תכנון ואופטימיזציה של טיולים")
    doc.add_paragraph()

    # 15.12.1
    add_heading(doc, "15.12.1 מבט על — ארכיטקטura כללית", level=2)
    add_para(doc, "מערכת TripLink היא מערכת לתכנון ואופטימיזציה של טיולים בישראל. המערכת בנויה בשכבות (Layered Architecture), כאשר כל שכבה אחראית על תחום אחריות מוגדר, והתקשורת בין השכבות מתבצעת באמצעות ממשקים (Interfaces) ואובייקטי העברת נתונים (DTOs).")

    add_table(doc,
        ["שכבה", "יחידות עיקריות", "תפקיד כללי"],
        [
            ["לקוח (Client)", "אפליקציית Angular", "ממשק משתמש — שליחת בקשות HTTP ותצוגת תוצאות"],
            ["API (Presentation)", "Controllers", "קבלת בקשות HTTP, החזרת תגובות JSON"],
            ["לוגיקה עסקית", "Services, Optimizer Pipeline", "עיבוד נתונים, CRUD, הרצת אלגוריתם"],
            ["גישה לנתונים", "Repositories, TripContext", "שליפה ושמירה בבסיס הנתונים"],
            ["מודל (Models)", "Entities, DTOs, מודלי Optimizer", "ייצוג נתונים פנימיים וחיצוניים"],
            ["אחסון", "SQL Server (TripLink)", "אחסון קבוע של נתוני המערכת"],
            ["שירות חיצוני", "Google Maps Directions API", "חישוב זמני נסיעה בתחבורה ציבורית"],
        ])

    add_module_tree(doc)

    # 15.12.2 Controllers
    add_heading(doc, "15.12.2 פירוט יחידות המערכת", level=2)
    add_heading(doc, "15.12.2.1 שכבת ה-API — Controllers", level=3)

    controllers = [
        ("TripsController", "נקודת כניסה לפעולות טיולים — יצירה, שליפה, מחיקה, אופטימיזציה ושמירת מסלול.",
         "בקשות HTTP עם CreateTripDto, OptimizeRequestDto, רשימת מזהי יעדים", "משתמש (Angular)",
         "TripDto, OptimizeResultDto, קודי סטטוס HTTP", "משתמש", "TripService, IOptimizerService"),
        ("UsersController", "ניהול משתמשים — יצירה, שליפה ומחיקה.",
         "CreateUserDto, מזהה משתמש", "משתמש (Angular)", "UserDto", "משתמש", "UserService"),
        ("DestinationsController", "שליפת יעדי טיול לפי מזהה, אזור, רמת קושי ותחנות.",
         "מזהה יעד, שם אזור, מזהה רמת קושי", "משתמש (Angular)", "DestinationDto, StationDto", "משתמש", "DestinationService"),
        ("LookupsController", "נתוני עזר לטופס — קטגוריות, רמות, סוגי מטיילים, features, אזורים.",
         "בקשות GET", "משתמש (Angular)", "CategoryDto, DifficultyLevelDto ועוד", "משתמש", "LookupService"),
    ]

    for name, role, inp, src, out, dest, deps in controllers:
        add_heading(doc, name, level=4)
        add_bullet(doc, f"תפקיד: {role}")
        add_bullet(doc, f"קלטים: {inp}")
        add_bullet(doc, f"מקור הקלטים: {src}")
        add_bullet(doc, f"פלטים: {out}")
        add_bullet(doc, f"יעד הפלטים: {dest}")
        add_bullet(doc, f"תלויות: {deps}")

    # Services
    add_heading(doc, "15.12.2.2 שכבת הלוגיקה העסקית — Services", level=3)
    services = [
        ("TripService", "ניהול מחזור חיים של טיול — CRUD, שיוך קטגוריות/תכונות, שמירת מסלול.",
         "CreateTripDto, מזהי טיול", "TripsController + SQL Server", "TripDto, כתיבה ל-DB", "Controller / DB", "TripContext"),
        ("UserService", "יצירה, שליפה ומחיקה של משתמשים.", "CreateUserDto", "UsersController + DB", "UserDto", "Controller / DB", "TripContext"),
        ("DestinationService", "שליפת יעדים ותחנות.", "מזהה, אזור, רמה", "DestinationsController + DB", "DestinationDto, StationDto", "Controller", "TripContext"),
        ("LookupService", "מתווך לנתוני lookup.", "—", "LookupsController", "DTOs של lookup", "Controller", "ILookupRepository"),
    ]
    for name, role, inp, src, out, dest, deps in services:
        add_heading(doc, name, level=4)
        add_bullet(doc, f"תפקיד: {role}")
        add_bullet(doc, f"קלטים: {inp} | מקור: {src}")
        add_bullet(doc, f"פלטים: {out} | יעד: {dest}")
        add_bullet(doc, f"תלויות: {deps}")

    # Optimizer
    add_heading(doc, "15.12.2.3 מנוע האופטימיזציה (Optimizer)", level=3)
    add_para(doc, "OptimizerServiceImpl מפעיל את OptimizerPipeline, המריץ 7 שלבים (Steps 0–6) על OptimizationContext משותף. OptimizeResultMapper ממיר את התוצאה ל-OptimizeResultDto.")

    add_table(doc,
        ["שלב", "שם", "תפקיד", "מקור קלט", "פלט"],
        [
            ["0", "InputLoader", "טעינת טיול ויעדים מה-DB", "OptimizerDataRepository", "OptimizerDestination, OptimizerParams"],
            ["1", "WeightCalculator", "חישוב משקלים", "Step 0", "WeightedDestination"],
            ["2", "ScoreTableBuilder", "טבלת עלויות מעבר", "Steps 0–1, Google Maps", "ScoreTable"],
            ["3", "DestinationRanker", "דירוג יעדים", "Steps 0–1", "רשימה מדורגת"],
            ["4", "InitialRouteBuilder", "מסלול התחלתי", "Steps 0–3", "InitialRoute"],
            ["5", "SaOptimizer", "Simulated Annealing", "Steps 0–4", "BestRoute"],
            ["6", "TripItineraryBuilder", "itinerary + narrative", "Steps 0–5, Repository", "TripPlan"],
        ])

    add_heading(doc, "ITransitApiService — שירות התחבורה", level=4)
    add_para(doc, "מספק זמני נסיעה בתחבורה ציבורית ובמכונית. מימוש: GoogleMapsTransitApiService (Google Maps API) עם MockTransitApiService כ-fallback.")

    # Data Access
    add_heading(doc, "15.12.2.4 שכבת גישה לנתונים", level=3)
    add_bullet(doc, "TripContext — DbContext מרכזי: DbSets, מפתחות מורכבים, קשרי FK.")
    add_bullet(doc, "LookupRepository — שליפת קטגוריות, רמות, סוגי מטייל, features, אזורים.")
    add_bullet(doc, "OptimizerDataRepository — טיול לאופטימיזציה, יעדים מסוננים, קווי אוטובוס.")

    add_heading(doc, "15.12.2.5 שכבת המודל (Models)", level=3)
    add_bullet(doc, "Entities (17) — ייצוג ORM של בסיס הנתונים.")
    add_bullet(doc, "DTOs (14) — תקשורת API ללא navigation properties.")
    add_bullet(doc, "מודלי Optimizer — OptimizationContext, OptimizerRoute, TripPlan (פנימיים).")

    add_heading(doc, "15.12.2.6 בסיס הנתונים — SQL Server", level=3)
    add_para(doc, "מסד TripLink מאחסן יעדים, תחנות, קווי אוטובוס, טיולים, משתמשים וטבלאות קשר. הסכמה מנוהלת בקבצי SQL (SeedData.sql, ClearData.sql).")

    add_heading(doc, "15.12.2.7 Google Maps API", level=3)
    add_para(doc, "שירות חיצוני לחישוב מסלולי תחבורה ציבורית. נקרא מ-ArcCostCalculator בשלב 2 של האופטימיזציה.")

    # Data flow
    add_heading(doc, "15.12.3 זרימת מידע במערכת", level=2)
    add_heading(doc, "15.12.3.1 זרימה כללית", level=3)
    add_para(doc, "המשתמש פועל באפליקציית Angular → Controller מקבל HTTP/JSON → Service מעבד → TripContext/Repository ניגש ל-SQL Server → DTO חוזר ל-Controller → JSON ל-Angular → תצוגה למשתמש.")

    add_heading(doc, "15.12.3.2 זרימת אופטימיזציה", level=3)
    steps = [
        "המשתמש שולח POST /api/trips/optimize עם OptimizeRequestDto.",
        "TripsController → OptimizerServiceImpl → OptimizerPipeline.",
        "Step 0: שליפת טיול ויעדים מ-OptimizerDataRepository (SQL Server).",
        "Step 1: חישוב משקלים לכל יעד.",
        "Step 2: בניית ScoreTable — ArcCostCalculator שולח שאילתות ל-Google Maps API.",
        "Step 3: דירוג יעדים | Step 4: מסלול התחלתי | Step 5: Simulated Annealing.",
        "Step 6: בניית TripPlan (זמנים, קווי אוטובוס, narrative בעברית).",
        "OptimizeResultMapper → OptimizeResultDto → JSON למשתמש.",
        "(אופציונלי) POST save-route → TripService → שמירה ב-Des_of_trip.",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"{i}. {s}")

    add_heading(doc, "15.12.3.3 זרימת יצירת טיול", level=3)
    add_bullet(doc, "Angular שולף lookups מ-LookupsController → LookupService → LookupRepository → DB.")
    add_bullet(doc, "POST /api/trips עם CreateTripDto → TripService → Trip + CategoriesToTrip + FeatureToTrip + NatureTrip → DB.")

    # Principles
    add_heading(doc, "15.12.4 עקרונות ארכיטקטוניים", level=2)
    principles = [
        "הפרדת אחריות (Separation of Concerns) — Controllers, Services, Repositories, Models.",
        "Pipeline Pattern — אופטימיזציה ב-7 שלבים (IOptimizerStep).",
        "Repository Pattern (חלקי) — Optimizer ו-Lookups דרך Repositories.",
        "Dependency Injection — רישום ב-Program.cs.",
        "DTO Pattern — API לא חושף Entities.",
        "Strategy / Fallback — ITransitApiService עם Google Maps ו-Mock.",
    ]
    for p in principles:
        add_bullet(doc, p)

    add_heading(doc, "15.12.5 סיכום", level=2)
    add_para(doc, "המערכת מאורגנת בשכבות: Angular → Controllers → Services/Optimizer → Repositories/TripContext → SQL Server, עם אינטegracja ל-Google Maps. זרימת המידע מתחילה בבקשת HTTP, עוברת דרך שכבות העיבוד, ומסתיימת בהחזרת DTO ללקוח או בשמירה בבסיס הנתונים.")

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_document()
