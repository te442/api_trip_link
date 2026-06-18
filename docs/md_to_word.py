"""Convert סעיף_21 markdown to Word with Hebrew RTL support."""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm

DOCS_DIR = Path(__file__).parent
MD_FILE = DOCS_DIR / "סעיף_21_קוד_התוכנית.md"
OUT_FILE = DOCS_DIR / "סעיף_21_קוד_התוכנית.docx"


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


def add_formatted_run(paragraph, text, bold=False, italic=False, code=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    else:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    return run


def parse_inline(paragraph, text):
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            add_formatted_run(paragraph, text[pos:m.start()])
        chunk = m.group(0)
        if chunk.startswith("**"):
            add_formatted_run(paragraph, chunk[2:-2], bold=True)
        elif chunk.startswith("`"):
            add_formatted_run(paragraph, chunk[1:-1], code=True)
        pos = m.end()
    if pos < len(text):
        add_formatted_run(paragraph, text[pos:])


def is_table_row(line):
    return line.strip().startswith("|") and "|" in line.strip()[1:]


def is_separator_row(line):
    s = line.strip()
    return s.startswith("|") and re.match(r"^\|[-:\s|]+\|$", s)


def add_table(doc, rows):
    if len(rows) < 2:
        return
    header = [c.strip() for c in rows[0].strip().strip("|").split("|")]
    data_rows = rows[2:] if len(rows) > 1 and is_separator_row("|".join(rows[1].split("|"))) else rows[1:]
    if len(rows) > 1 and all(c.strip().replace("-", "").replace(":", "") == "" for c in rows[1].split("|")):
        data_rows = rows[2:]
    elif len(rows) > 1 and re.match(r"^\|[-:\s|]+\|$", rows[1].strip()):
        data_rows = rows[2:]
    else:
        data_rows = rows[1:]

    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        set_rtl(p)
        add_formatted_run(p, h, bold=True)

    for row_line in data_rows:
        if not row_line.strip():
            continue
        cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
        if len(cells) < len(header):
            cells += [""] * (len(header) - len(cells))
        tr = table.add_row()
        for i, val in enumerate(cells[:len(header)]):
            tr.cells[i].text = ""
            p = tr.cells[i].paragraphs[0]
            set_rtl(p)
            parse_inline(p, val)

    doc.add_paragraph()


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    set_rtl(p)
    parse_inline(p, text)
    for run in p.runs:
        run.font.name = "Calibri"


def add_code_block(doc, lines, lang=""):
    p = doc.add_paragraph()
    set_rtl(p)
    p.paragraph_format.left_indent = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    label = f"[{lang}] " if lang else ""
    full = label + "\n".join(lines)
    run = p.add_run(full)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)


def convert(md_path: Path, out_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Default font for Normal style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines, lang)
            i += 1
            continue

        if is_table_row(line):
            table_rows = []
            while i < len(lines) and is_table_row(lines[i]):
                table_rows.append(lines[i])
                i += 1
            add_table(doc, table_rows)
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            add_heading(doc, text, min(level, 4))
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_rtl(p)
            parse_inline(p, stripped[2:])
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        p = doc.add_paragraph()
        set_rtl(p)
        parse_inline(p, stripped)
        i += 1

    doc.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    convert(MD_FILE, OUT_FILE)
