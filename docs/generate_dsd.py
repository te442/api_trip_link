# -*- coding: utf-8 -*-
"""Generates DSD (Database Schema Description) Word document for TripLink."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\User\Desktop\פרויקט טיול\API_trip_link\docs\DSD_תיאור_מסד_נתונים.docx"

COL_HEADERS = ["עמודה", "תפקיד", "טיפוס נתונים", "חובה", "הערות"]


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            set_rtl(p)
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                set_rtl(p)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_entity_section(doc, title, role, relations, columns):
    add_heading(doc, title, level=2)
    add_para(doc, f"תפקיד: {role}")
    if relations:
        add_para(doc, f"קשרים: {relations}")
    add_table(doc, COL_HEADERS, columns)


TABLES = [
    (
        "4.1 Users — משתמשים",
        "אחסון פרטי משתמשי המערכת ואימות (JWT).",
        "1:N → Trip",
        [
            ("user_id", "PK", "NVARCHAR", "כן", "מזהה ייחודי (למשל U001)"),
            ("FullName", "רגיל", "NVARCHAR", "כן", "שם מלא"),
            ("Phon", "רגיל", "NVARCHAR", "כן", "מספר טלפון"),
            ("Email", "רגיל", "NVARCHAR(256)", "לא", "דוא\"ל לכניסה והרשמה"),
            ("PasswordHash", "רגיל", "NVARCHAR(500)", "לא", "סיסמה מוצפנת (BCrypt)"),
        ],
    ),
    (
        "4.2 Trip — טיולים",
        "טיול מתוכנן שמשתמש יוצר — תאריך, שעות, נקודת התחלה.",
        "N:1 ← Users; 1:N → Nature_trip; M:N ↔ Destination, Categories, FeatureTypes",
        [
            ("Trip_id", "PK", "INT (IDENTITY)", "כן", "מזהה טיול"),
            ("Trip_name", "רגיל", "NVARCHAR", "כן", "שם הטיול"),
            ("user_id", "FK → Users", "NVARCHAR", "כן", "בעל הטיול"),
            ("Trip_Date", "רגיל", "DATETIME", "לא", "תאריך הטיול"),
            ("Address_start", "רגיל", "NVARCHAR", "כן", "כתובת / נקודת יציאה"),
            ("Start_time", "רגיל", "TIME", "לא", "שעת יציאה (למשל 08:00)"),
            ("End_time", "רגיל", "TIME", "לא", "שעת חזרה (למשל 18:00)"),
            ("trip_cost", "רגיל", "DECIMAL", "לא", "עלות משוערת; ברירת מחדל 0"),
        ],
    ),
    (
        "4.3 Nature_trip — העדפות טיול",
        "פרמטרי תכנון לטיול — אזור, מינימום/מקסימום יעדים, רמת קושי.",
        "N:1 → Trip, Difficulty_level",
        [
            ("natureTrip_id", "PK", "INT (IDENTITY)", "כן", "מזהה רשומה"),
            ("trip_id", "FK → Trip", "INT", "כן", "טיול משויך"),
            ("Num_break", "רגיל", "INT", "לא", "מספר הפסקות"),
            ("Min_num_des", "רגיל", "INT", "לא", "מינימום יעדים"),
            ("Max_num_des", "רגיל", "INT", "לא", "מקסימום יעדים"),
            ("level_id", "FK → Difficulty_level", "INT", "לא", "רמת קושי מבוקשת"),
            ("Region", "רגיל", "NVARCHAR(50)", "לא", "אזור טיול (דרום/צפון/מרכז)"),
        ],
    ),
    (
        "4.4 Des_of_trip — יעדים במסלול (טבלת קישור)",
        "קישור M:N בין Trip ל-Destination; שמירת מסלול לאחר אופטימיזציה.",
        "M:N בין Trip ו-Destination",
        [
            ("Trip_id", "PK, FK → Trip", "INT", "כן", "חלק ממפתח מורכב"),
            ("Des_id", "PK, FK → Destination", "INT", "כן", "חלק ממפתח מורכב"),
            ("visit_number", "רגיל", "INT", "לא", "סדר ביקור ביעד"),
        ],
    ),
    (
        "4.5 Categories — קטגוריות",
        "טבלת עזר — סיווג יעדים וטיולים (טבע, היסטוריה, חוף וכו').",
        "M:N ↔ Destination, M:N ↔ Trip",
        [
            ("categories_id", "PK", "INT (IDENTITY)", "כן", "מזהה קטגוריה"),
            ("categories_name", "רגיל", "NVARCHAR", "כן", "שם הקטגוריה"),
        ],
    ),
    (
        "4.6 Categories_to_trip — קטגוריות לטיול (טבלת קישור)",
        "קטגוריות שהמשתמש בחר בשלב 2 של אשף התכנון.",
        "M:N בין Categories ו-Trip",
        [
            ("categories_id", "PK, FK → Categories", "INT", "כן", "חלק ממפתח מורכב"),
            ("trip_id", "PK, FK → Trip", "INT", "כן", "חלק ממפתח מורכב"),
        ],
    ),
    (
        "4.7 Categories_of_Destination — קטגוריות ליעד (טבלת קישור)",
        "לאיזה קטגוריות שייך כל יעד (לא לטיול).",
        "M:N בין Categories ו-Destination",
        [
            ("Categories_id", "PK, FK → Categories", "INT", "כן", "חלק ממפתח מורכב"),
            ("Des_id", "PK, FK → Destination", "INT", "כן", "חלק ממפתח מורכב"),
        ],
    ),
    (
        "4.8 FeatureTypes — סוגי מאפיינים",
        "מאפיינים נדרשים (חניה חינם, נגישות, כניסה חינם וכו').",
        "M:N ↔ Trip",
        [
            ("Feature_id", "PK", "INT (IDENTITY)", "כן", "מזהה מאפיין"),
            ("Feature", "רגיל", "NVARCHAR", "כן", "תיאור המאפיין"),
        ],
    ),
    (
        "4.9 Feature_to_trip — מאפיינים לטיול (טבלת קישור)",
        "מאפיינים שהמשתמש דרש לטיול.",
        "M:N בין FeatureTypes ו-Trip",
        [
            ("Feature_id", "PK, FK → FeatureTypes", "INT", "כן", "במודל C# מסומן PK בלבד"),
            ("trip_id", "FK → Trip", "INT", "כן", "בפועל מפתח מורכב עם Feature_id"),
        ],
    ),
    (
        "4.10 Destination — יעדי טיול",
        "אתרי טיול — שם, אזור, קושי, זמן ביקור, מיקום גיאוגרפי.",
        "N:1 → Difficulty_level, TypeTraveler; M:N ↔ Categories, Station, Trip",
        [
            ("Des_id", "PK", "INT (IDENTITY)", "כן", "מזהה יעד"),
            ("Name_des", "רגיל", "NVARCHAR", "כן", "שם היעד"),
            ("Rregion", "רגיל", "NVARCHAR", "כן", "אזור גיאוגרפי"),
            ("level_id", "FK → Difficulty_level", "INT", "לא", "רמת קושי"),
            ("Traveler_id", "FK → TypeTraveler", "INT", "לא", "סוג מטייל מומלץ"),
            ("Time_des", "רגיל", "TIME", "לא", "משך ביקור מומלץ"),
            ("lat", "רגיל", "DECIMAL(9,6)", "לא", "קו רוחב (למפה)"),
            ("lon", "רגיל", "DECIMAL(9,6)", "לא", "קו אורך (למפה)"),
            ("image_url", "רגיל", "NVARCHAR(500)", "לא", "קישור לתמונת יעד"),
        ],
    ),
    (
        "4.11 Difficulty_level — רמות קושי",
        "טבלת עזר — קל / בינוני / קשה.",
        "1:N → Destination, Nature_trip",
        [
            ("level_id", "PK", "INT", "כן", "מזהה רמה"),
            ("level_type", "רגיל", "NVARCHAR", "כן", "שם הרמה"),
        ],
    ),
    (
        "4.12 TypeTraveler — סוגי מטיילים",
        "טבלת עזר — משפחות, זוגות, יחידים, קבוצות.",
        "1:N → Destination",
        [
            ("Traveler_id", "PK", "INT", "כן", "מזהה סוג"),
            ("TypeTraveler", "רגיל", "NVARCHAR", "כן", "שם הסוג"),
        ],
    ),
    (
        "4.13 DestinationFeatures — מאפייני יעד",
        "טבלה עזרית; קישור ל-Destination לא מוגדר ב-Entity Framework.",
        "אין FK מוגדר ב-TripContext",
        [
            ("Feature_id", "PK", "INT", "כן", "מזהה מאפיין"),
            ("topical_id", "רגיל", "INT", "כן", "מזהה נושאי"),
        ],
    ),
    (
        "4.14 Agency — חברות אוטובוס",
        "מפעילי תחבורה (אגד, דן, נתיב אקספרס וכו'). קיימת ב-SQL; לא ממופה ב-TripContext.",
        "1:N → Bus",
        [
            ("agency_id", "PK", "INT (IDENTITY)", "כן", "מזהה חברה"),
            ("external_agency_id", "רגיל", "INT", "כן", "מזהה חיצוני (ממשלתי)"),
            ("agency_name", "רגיל", "NVARCHAR", "כן", "שם החברה"),
            ("phon", "רגיל", "NVARCHAR", "כן", "טלפון"),
        ],
    ),
    (
        "4.15 Bus — קווי אוטובוס",
        "קווי תחבורה ציבורית — מספר קו, כיוון, חברה מפעילה.",
        "N:1 → Agency; M:N ↔ Station",
        [
            ("bus_id", "PK", "INT (IDENTITY)", "כן", "מזהה קו"),
            ("Bus_code", "רגיל", "NVARCHAR", "כן", "קוד פנימי (B-060)"),
            ("Bus_number", "רגיל", "INT", "כן", "מספר קו (60, 384)"),
            ("agency_id", "FK → Agency", "INT", "לא", "חברת ההפעלה"),
            ("Direction", "רגיל", "NVARCHAR", "כן", "כיוון הנסיעה"),
            ("government_route_id", "רגיל", "NVARCHAR", "כן", "מזהה מסלול ממשלתי"),
        ],
    ),
    (
        "4.16 Station — תחנות אוטובוס",
        "תחנות עצירה — שם, אזור, קואורדינטות.",
        "M:N ↔ Bus, M:N ↔ Destination",
        [
            ("Station_num", "PK", "INT", "כן", "מזהה תחנה"),
            ("Statoin_code", "רגיל", "NVARCHAR", "כן", "קוד תחנה"),
            ("Station_name", "רגיל", "NVARCHAR", "כן", "שם התחנה"),
            ("area", "רגיל", "NVARCHAR", "כן", "אזור (דרום/צפון/מרכז)"),
            ("lat", "רגיל", "DECIMAL", "לא", "קו רוחב"),
            ("lon", "רגיל", "DECIMAL", "לא", "קו אורך"),
            ("government_stop_id", "רגיל", "NVARCHAR", "כן", "מזהה תחנה ממשלתי"),
        ],
    ),
    (
        "4.17 bus_station — תחנות בקו (טבלת קישור)",
        "אילו תחנות עוברים בכל קו וסדר העצירה.",
        "M:N בין Bus ו-Station",
        [
            ("bus_id", "PK, FK → Bus", "INT", "כן", "חלק ממפתח מורכב"),
            ("station_id", "PK, FK → Station", "INT", "כן", "מצביע ל-Station_num"),
            ("stop_sequence", "רגיל", "INT", "לא", "סדר עצירה במסלול"),
        ],
    ),
    (
        "4.18 station_to_destination — נגישות יעד מתחנה (טבלת קישור)",
        "תחנת הירידה הקרובה ליעד, מרחק והליכה.",
        "M:N בין Destination ו-Station",
        [
            ("Des_id", "PK, FK → Destination", "INT", "כן", "חלק ממפתח מורכב"),
            ("Station_num", "PK, FK → Station", "INT", "כן", "חלק ממפתח מורכב"),
            ("Direction_Type", "רגיל", "NVARCHAR", "כן", "כיוון (הלוך/חזור)"),
            ("Walking_distance", "רגיל", "FLOAT", "לא", "מרחק הליכה (ק\"מ)"),
            ("Walking time", "רגיל", "TIME", "לא", "זמן הליכה"),
            ("Walking instructions", "רגיל", "NVARCHAR", "כן", "הוראות הליכה"),
            ("level_id", "רגיל", "INT", "לא", "רמת קושי הליכה"),
            ("Feature_id", "רגיל", "INT", "לא", "מאפיין נלווה"),
        ],
    ),
]


def build():
    doc = Document()

    add_heading(doc, "תיאור מסד הנתונים (DSD)", level=1)
    add_para(doc, "מערכת TripLink — תכנון ואופטימיזציה של טיולים")
    doc.add_paragraph()

    # 1. Overview
    add_heading(doc, "1. סקירה כללית", level=2)
    add_para(doc,
        "מסד הנתונים TripLink (Microsoft SQL Server) תומך במערכת לתכנון ואופטימיזציה "
        "של טיולים יומיים בתחבורה ציבורית. הנתונים מאורגנים בארבעה תחומים לוגיים:")
    add_table(doc, ["תחום", "טבלאות עיקריות", "תפקיד"],
        [
            ("משתמשים וטיולים", "Users, Trip, Nature_trip, Des_of_trip, Categories_to_trip, Feature_to_trip",
             "ניהול משתמשים, יצירת טיולים, העדפות ומסלול"),
            ("יעדים ומאפיינים", "Destination, Difficulty_level, TypeTraveler, Categories, Categories_of_Destination, FeatureTypes",
             "קטלוג יעדי טיול, סיווג ומאפיינים"),
            ("תחבורה ציבורית", "Agency, Bus, Station, bus_station, station_to_destination",
             "קווי אוטובוס, תחנות ונגישות ליעדים"),
            ("טבלאות קישור", "6 טבלאות M:N", "מממשות קשרי רבים-לרבים"),
        ])
    add_para(doc, "סה\"כ: 18 טבלאות (17 במיפוי Entity Framework + Agency ב-SQL בלבד).")
    add_para(doc, "גישה למסד: Entity Framework Core (TripContext) — שאילתות LINQ דרך Repositories ו-Services.")

    # 2. ERD
    add_heading(doc, "2. תרשים ERD", level=2)
    add_para(doc, "תרשים ישויות-קשרים (ERD) מלא מוצג בנספח הגרפי (קובץ erd_visual.html / צילום מסך).")
    add_para(doc, "התרשים מציג את כל הישויות, מפתחות ראשיים וזרים, וקשרי 1:N ו-M:N בין הטבלאות.")

    # 3. Relationships
    add_heading(doc, "3. סיכום קשרים בין טבלאות", level=2)
    add_heading(doc, "3.1 קשרי 1:N (אחד לרבים)", level=3)
    add_table(doc, ["מישות מקור", "מישות יעד", "עמודת FK", "דוגמה"],
        [
            ("Users", "Trip", "user_id", "משתמש אחד → הרבה טיולים"),
            ("Trip", "Nature_trip", "trip_id", "טיול אחד → העדפות טבע"),
            ("Difficulty_level", "Destination", "level_id", "רמת קושי → יעדים רבים"),
            ("Difficulty_level", "Nature_trip", "level_id", "רמת קושי → טיולים רבים"),
            ("TypeTraveler", "Destination", "Traveler_id", "סוג מטייל → יעדים רבים"),
            ("Agency", "Bus", "agency_id", "אגד → קווים 60, 384, 392"),
        ])
    add_heading(doc, "3.2 קשרי M:N (רבים לרבים)", level=3)
    add_table(doc, ["טבלה א", "טבלת קישור", "טבלה ב", "דוגמה"],
        [
            ("Trip", "Des_of_trip", "Destination", "טיול ↔ יעדים במסלול"),
            ("Trip", "Categories_to_trip", "Categories", "קטגוריות שנבחרו לטיול"),
            ("Trip", "Feature_to_trip", "FeatureTypes", "מאפיינים נדרשים בטיול"),
            ("Destination", "Categories_of_Destination", "Categories", "קטגוריות של כל יעד"),
            ("Bus", "bus_station", "Station", "קו עוצר בתחנות"),
            ("Destination", "station_to_destination", "Station", "יעד נגיש מתחנה"),
        ])

    # 4. Tables detail
    add_heading(doc, "4. תיאור מפורט לכל טבלה", level=2)
    add_para(doc, "לכל טבלה מפורטים: שם העמודה, תפקיד (PK/FK/רגיל), טיפוס נתונים, האם חובה (Allow Null), והערות.")

    for title, role, relations, columns in TABLES:
        add_entity_section(doc, title, role, relations, columns)

    # 5. Views / SP
    add_heading(doc, "5. Views, Stored Procedures ואובייקטי מסד נוספים", level=2)
    add_table(doc, ["סוג אובייקט", "שימוש בפרויקט"],
        [
            ("Views", "לא נעשה שימוש"),
            ("Stored Procedures", "לא נעשה שימוש"),
            ("Triggers", "לא נעשה שימוש"),
            ("Functions (UDF)", "לא נעשה שימוש"),
        ])
    add_para(doc,
        "כל הגישה למסד הנתונים מתבצעת דרך Entity Framework Core. "
        "שינויי סכימה בוצעו באמצעות סקריפטי ALTER: AlterUsersAuth.sql, AlterDestination.sql, AlterNatureTrip.sql.")

    # 6. Notes
    add_heading(doc, "6. הערות טכניות", level=2)
    notes = [
        "שתי טבלאות קטגוריות שונות: Categories_to_trip (קישור לטיול) לעומת Categories_of_Destination (קישור ליעד).",
        "טבלת Agency קיימת ב-SQL Server אך לא רשומה ב-TripContext — הקשר ל-Bus הוא לוגי בלבד.",
        "טבלת DestinationFeatures — DbSet קיים ב-EF ללא הגדרת FK ל-Destination.",
        "טבלת Feature_to_trip — במודל C# רק Feature_id מסומן כ-PK; במבנה הטבלה המפתח מורכב (Feature_id, trip_id).",
        "שישה מפתחות מורכבים בטבלאות קישור M:N.",
        "אין שימוש ב-Views או Stored Procedures — ארכיטקטורת מסד נתונים פשוטה ושקופה.",
    ]
    for n in notes:
        add_bullet(doc, n)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
