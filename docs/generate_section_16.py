# -*- coding: utf-8 -*-
"""Generates section 16 Word document — Screen Flow Diagram for TripLink."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\User\Desktop\פרויקט טיול\API_trip_link\docs\16_תרשים_זרימת_מסכים.docx"


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            set_rtl(p)
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                set_rtl(p)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
    doc.add_paragraph()


def add_screen_tree(doc):
    add_heading(doc, "16.2.1 היררכיית המסכים (Screen Hierarchy)", level=3)
    tree = [
        "מתכנן טיולים — TripLink (Angular SPA)",
        "│",
        "├── אזור ציבורי (ללא התחברות)",
        "│   ├── מסך התחברות (/login)",
        "│   └── מסך הרשמה (/register)",
        "│",
        "├── מסך מעטפת (App Shell) — תפריט ניווט (מוצג לאחר התחברות)",
        "│   ├── קישור: הטיולים שלי",
        "│   ├── קישור: + תכנן טיול",
        "│   ├── קישור: יעדים",
        "│   └── כפתור: התנתק",
        "│",
        "├── מסך 1: הטיולים שלי (/my-trips)",
        "│   ├── טבלה: טיולים של המשתמש המחובר",
        "│   └── פעולות: צפה במסלול | מחק | תכנן טיול חדש",
        "│",
        "├── מסך 2: אשף תכנון טיול (/plan) — 3 שלבים",
        "│   ├── שלב 1: פרטי בסיס (שם, אזור, תאריך, כתובת)",
        "│   ├── שלב 2: העדפות (קושי, קטגוריות, מאפיינים, מינ/מקס יעדים)",
        "│   └── שלב 3: זמנים ופרמטרי אופטימיזציה",
        "│",
        "├── מסך 3: חישוב מסלול (/plan/optimize/:tripId)",
        "│   ├── סיכום הטיול שנוצר",
        "│   └── כפתור: חשב מסלול אופטימלי",
        "│",
        "├── מסך 4: תוצאות הטיול (/trips/:id/result)  ← מסך מרכזי",
        "│   ├── מפת Google Maps + שרטוט מסלול",
        "│   ├── Timeline: יעד אחר יעד",
        "│   ├── כרטיסי יעדים (תמונה, זמנים, תחבורה)",
        "│   └── סיכום מילולי (Narrative)",
        "│",
        "└── מסך 5: רשימת יעדים (/destinations)",
        "    └── טבלה: כל היעדים במערכת (תצוגה בלבד)",
    ]
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run("\n".join(tree))
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def add_flow_diagram(doc):
    add_heading(doc, "16.2.2 תרשים מעברים בין מסכים (Screen Flow)", level=3)
    add_para(doc, "התרשים הבא מתאר את המעברים העיקריים בין המסכים. חצים מייצגים מעבר שמתבצע בפועל בקוד האפליקציה.")

    flow = [
        "                 ┌─────────────────────┐",
        "                 │  כניסה לאפליקציה   │",
        "                 │      (URL: /)       │",
        "                 └──────────┬──────────┘",
        "                            │ redirect",
        "                            ▼",
        "                 ┌─────────────────────┐",
        "                 │    התחברות        │",
        "                 │    (/login)       │",
        "                 └────┬────────┬─────┘",
        "                      │        │",
        "           הרשמה      │        │  התחברות (הצלחה)",
        "                      ▼        ▼",
        "           ┌──────────────┐  ┌──────────────────┐",
        "           │   הרשמה      │  │  הטיולים שלי     │",
        "           │ (/register)  │  │  (/my-trips)     │",
        "           └──────┬───────┘  └────────┬─────────┘",
        "                  │ הצלחה             │",
        "                  └──────────────────►│",
        "                                        │",
        "                          + תכנן טיול  │",
        "                                        ▼",
        "                          ┌────────────────────────┐",
        "                          │  אשף תכנון טיול        │",
        "                          │  (/plan) — 3 שלבים    │",
        "                          └────────────┬───────────┘",
        "                                       │ צור טיול (הצלחה)",
        "                                       ▼",
        "                          ┌────────────────────────┐",
        "                          │  חישוב מסלול           │",
        "                          │  (/plan/optimize/:id)  │",
        "                          └────────────┬───────────┘",
        "                                       │ חשב מסלול (הצלחה)",
        "                                       ▼",
        "                          ┌────────────────────────┐",
        "                          │  תוצאות הטיול          │",
        "                          │  (/trips/:id/result)   │",
        "                          │  מפה + פירוט מסלול    │",
        "                          └────────────┬───────────┘",
        "                                       │",
        "              ┌─────────────────────────┼─────────────────┐",
        "              ▼                         ▼                 ▼",
        "   ┌──────────────────┐      ┌─────────────────┐  (תפריט עליון",
        "   │  הטיולים שלי    │      │  רשימת יעדים    │   — מעבר חופשי",
        "   │  (/my-trips)     │      │ (/destinations) │   בין מסכים)",
        "   └──────────────────┘      └─────────────────┘",
    ]
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run("\n".join(flow))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_heading("ספר פרויקט גמר — TripLink", level=0)
    set_rtl(title)
    add_heading(doc, "סעיף 16 — תרשים זרימת מסכים (Screen Flow Diagram)", level=1)
    add_para(doc, "מקצוע: מדעי המחשב | רכיב: אפליקציית Angular — API_trip_link/trip-planner-app")
    doc.add_paragraph()

    # 16.1
    add_heading(doc, "16.1 מבוא", level=2)
    add_para(doc, "פרק זה מתאר את מבנה המסכים באפליקציית הלקוח (Frontend) של מערכת TripLink, את ההיררכיה ביניהם ואת המעברים (Navigation Flow) שמתרחשים במהלך שימוש במערכת. האפליקציה בנויה כ-Single Page Application (SPA) ב-Angular 19, עם ניווט מבוסס Router, אימות JWT ו-AuthGuard על מסכים פנימיים.")
    add_para(doc, "מערכת TripLink תומכת בתהליך מלא: התחברות → תכנון טיול (אשף 3 שלבים) → חישוב מסלול אופטימלי → הצגת תוצאות עם מפת Google Maps, פירוט אוטובוסים, תחנות ויעדים עם תמונות.")

    # 16.2
    add_heading(doc, "16.2 מבנה המסכים והיררכיה", level=2)
    add_para(doc, "האפליקציה כוללת שני מסכי אימות (התחברות והרשמה), מסך מעטפת (Shell), וחמישה מסכי תוכן עיקריים. מסכים פנימיים מוגנים על ידי AuthGuard — משתמש לא מחובר מופנה אוטומטית ל-/login.")

    add_table(doc,
        ["#", "שם המסך", "נתיב (Route)", "Component", "תפקיד"],
        [
            ["—", "מסך מעטפת", "—", "AppComponent", "תפריט ניווט + router-outlet (לאחר התחברות)"],
            ["0א", "התחברות", "/login", "LoginComponent", "כניסה עם אימייל וסיסמה (JWT)"],
            ["0ב", "הרשמה", "/register", "RegisterComponent", "יצירת משתמש חדש"],
            ["1", "הטיולים שלי", "/my-trips", "MyTripsComponent", "רשימת טיולים של המשתמש"],
            ["2", "אשף תכנון טיול", "/plan", "TripWizardComponent", "תכנון טיול ב-3 שלבים"],
            ["3", "חישוב מסלול", "/plan/optimize/:tripId", "OptimizeScreenComponent", "הרצת אלגוריתם אופטימיזציה"],
            ["4", "תוצאות הטיול", "/trips/:id/result", "TripResultComponent", "מפה, timeline, פירוט מסלול"],
            ["5", "רשימת יעדים", "/destinations", "DestinationsListComponent", "צפייה בכל היעדים"],
        ])

    add_screen_tree(doc)
    add_flow_diagram(doc)

    # 16.3 Screen details
    add_heading(doc, "16.3 פירוט מסכים ומעברים", level=2)

    screens = [
        ("16.3.1 מסך התחברות (/login)", [
            "מסך ברירת המחדל — כניסה ל-URL / מפנה אוטומטית ל-/login.",
            "קלטים: אימייל, סיסמה.",
            "API: POST /api/auth/login → מחזיר JWT.",
            "מעבר «התחבר» (הצלחה): → /my-trips.",
            "מעבר «הירשם כאן»: → /register.",
            "אם המשתמש כבר מחובר: redirect אוטומטי ל-/my-trips.",
        ]),
        ("16.3.2 מסך הרשמה (/register)", [
            "קלטים: שם מלא, אימייל, טלפון, סיסמה.",
            "API: POST /api/auth/register → מחזיר JWT.",
            "מעבר «הירשם» (הצלחה): → /my-trips.",
            "מעבר «התחבר כאן»: → /login.",
        ]),
        ("16.3.3 מסך מעטפת (App Shell)", [
            "מוצג רק למשתמש מחובר (auth.isLoggedIn()).",
            "תפריט: הטיולים שלי | + תכנן טיול | יעדים | התנתק.",
            "מציג את המסך הפעיל באמצעות router-outlet.",
            "«התנתק»: מוחק JWT מ-localStorage → /login.",
        ]),
        ("16.3.4 מסך הטיולים שלי (/my-trips)", [
            "API: GET /api/trips/user/{userId} — טיולים של המשתמש המחובר.",
            "פעולות: «צפה במסלול» → /trips/{id}/result.",
            "פעולות: «מחק» → confirm → DELETE /api/trips/{id}.",
            "«+ תכנן טיול חדש»: → /plan.",
        ]),
        ("16.3.5 אשף תכנון טיול (/plan) — 3 שלבים", [
            "שלב 1 — פרטי בסיס: שם טיול, אזור, תאריך, כתובת התחלה.",
            "שלב 2 — העדפות: רמת קושי, קטגוריות (multi-select), מאפיינים, מינ/מקס יעדים.",
            "שלב 3 — זמנים: שעת יציאה/חזרה, זמן נסיעה מקסימלי, זמן חזרה, יעילות תחבורה.",
            "Lookups נטענים מ-/api/lookups/* (קטגוריות, קושי, מאפיינים, אזורים).",
            "«צור טיול והמשך לאופטימיזציה»: POST /api/trips (עם userId מה-JWT) → /plan/optimize/{tripId}.",
            "«ביטול»: → /my-trips.",
        ]),
        ("16.3.6 מסך חישוב מסלול (/plan/optimize/:tripId)", [
            "מציג סיכום read-only של הטיול שנוצר.",
            "פרמטרי זמן מועברים ב-queryParams מהאשף.",
            "«חשב מסלול אופטימלי»: POST /api/trips/optimize.",
            "בהצלחה: POST /api/trips/{id}/save-route → שמירת מסלול → /trips/{id}/result.",
            "תוצאות נשמרות ב-sessionStorage (TripStateService) לצורך הצגה מיידית.",
        ]),
        ("16.3.7 מסך תוצאות הטיול (/trips/:id/result)", [
            "מסך התוצאה הסופי — המרכזי במערכת.",
            "מפת Google Maps: markers ממוספרים + polyline של המסלול.",
            "Timeline: לכל יעד — תמונה, שעות הגעה/עזיבה, משך שהייה.",
            "פירוט תחבורה: קווי אוטובוס, תחנות עלייה/ירידה, זמני נסיעה, הליכה.",
            "סיכום מילולי (Narrative) מה-API.",
            "מקור נתונים: TripStateService (מיידי) או GET /api/trips/{id}/itinerary.",
            "«חזור לטיולים שלי»: → /my-trips | «תכנן טיול חדש»: → /plan.",
        ]),
        ("16.3.8 מסך רשימת יעדים (/destinations)", [
            "מציג טבלת יעדים (GET /api/destinations) — שם, אזור, קושי, סוג מטייל, זמן ביקור.",
            "מסך תצוגה בלבד — מעברים רק דרך תפריט הניווט.",
        ]),
    ]

    for title, bullets in screens:
        add_heading(doc, title, level=3)
        for b in bullets:
            add_bullet(doc, b)

    # 16.4 Transitions table
    add_heading(doc, "16.4 טבלת מעברים בין מסכים", level=2)
    add_table(doc,
        ["מסך מקור", "פעולת משתמש / אירוע", "מסך יעד", "סוג מעבר"],
        [
            ["כניסה (/)", "Redirect אוטומטי", "/login", "Router redirect"],
            ["התחברות", "התחברות מוצלחת", "/my-trips", "Router.navigate"],
            ["התחברות", "לחיצה «הירשם כאן»", "/register", "RouterLink"],
            ["הרשמה", "הרשמה מוצלחת", "/my-trips", "Router.navigate"],
            ["הרשמה", "לחיצה «התחבר כאן»", "/login", "RouterLink"],
            ["תפריט עליון", "לחיצה «הטיולים שלי»", "/my-trips", "Navigation"],
            ["תפריט עליון", "לחיצה «+ תכנן טיול»", "/plan", "Navigation"],
            ["תפריט עליון", "לחיצה «יעדים»", "/destinations", "Navigation"],
            ["תפריט עליון", "לחיצה «התנתק»", "/login", "AuthService.logout"],
            ["הטיולים שלי", "לחיצה «+ תכנן טיול חדש»", "/plan", "RouterLink"],
            ["הטיולים שלי", "לחיצה «צפה במסלול»", "/trips/{id}/result", "RouterLink"],
            ["הטיולים שלי", "לחיצה «מחק» + אישור", "/my-trips (אותו מסך)", "רענון נתונים"],
            ["אשף תכנון", "לחיצה «הבא» (שלבים 1→2→3)", "/plan (אותו מסך)", "שינוי step"],
            ["אשף תכנון", "«צור טיול» (הצלחה)", "/plan/optimize/{tripId}", "Router.navigate"],
            ["אשף תכנון", "לחיצה «ביטול»", "/my-trips", "RouterLink"],
            ["חישוב מסלול", "«חשב מסלול» (הצלחה)", "/trips/{id}/result", "Router.navigate"],
            ["תוצאות הטיול", "«חזור לטיולים שלי»", "/my-trips", "RouterLink"],
            ["תוצאות הטיול", "«תכנן טיול חדש»", "/plan", "RouterLink"],
            ["AuthGuard", "גישה ללא JWT", "/login", "Guard redirect"],
        ])

    # 16.5 Main user flow
    add_heading(doc, "16.5 זרימת משתמש עיקרית (Main User Flow)", level=2)
    add_para(doc, "זרימת העבודה המרכזית של המערכת — מסלול השימוש הטיפוסי:")

    main_flow = [
        "1. המשתמש נכנס לאפליקציה → מועבר למסך «התחברות».",
        "2. משתמש חדש נרשם (/register) או משתמש קיים מתחבר (/login) → JWT נשמר.",
        "3. מעבר אוטומטי ל«הטיולים שלי» — רשימת הטיולים של המשתמש.",
        "4. לחיצה «+ תכנן טיול» → אשף 3 שלבים: פרטים → העדפות → זמנים.",
        "5. שלב 2 טוען Lookups מה-API (קטגוריות, אזורים, רמות קושי, מאפיינים).",
        "6. «צור טיול» → POST /api/trips → מעבר ל«חישוב מסלול».",
        "7. «חשב מסלול אופטימלי» → POST /api/trips/optimize → Pipeline אופטימיזציה.",
        "8. שמירת מסלול → מעבר ל«תוצאות הטיול».",
        "9. מסך התוצאות: מפת Google Maps, timeline, תמונות יעדים, פירוט אוטובוסים ותחנות.",
        "10. המשתמש יכול לחזור ל«הטיולים שלי» או לתכנן טיול חדש.",
    ]
    for step in main_flow:
        add_bullet(doc, step)

    add_heading(doc, "16.6 אבטחה וניווט", level=2)
    add_bullet(doc, "AuthGuard (authGuard) — מגן על כל המסכים הפנימיים; מפנה ל-/login אם אין JWT.")
    add_bullet(doc, "authInterceptor — מוסיף Authorization: Bearer {token} לכל בקשת HTTP ל-API.")
    add_bullet(doc, "AuthService — שומר JWT ופרטי משתמש ב-localStorage.")

    add_heading(doc, "16.7 סיכום", level=2)
    add_para(doc, "אפליקציית TripLink כוללת 8 מסכים (כולל אימות) ומסך מעטפת, המחוברים באמצעות Angular Router. הזרימה המרכזית: התחברות → תכנון (אשף) → אופטימיזציה → תוצאות עם מפה. מסך «תוצאות הטיול» הוא נקודת הסיום והחשובה ביותר — מציג למשתמש את המסלול המלא בצורה ויזואלית ומפורטת. תרשימי הזרימה והטבלאות בפרק זה משקפים את המימוש בקוד ב-trip-planner-app.")

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_document()
