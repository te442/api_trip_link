# -*- coding: utf-8 -*-
"""Generates User Guide + Testing & Evaluation Word document for TripLink."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"c:\Users\User\Desktop\פרויקט טיול\API_trip_link\docs\מדריך_משתמש_ובדיקות.docx"


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    set_rtl(h)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            set_rtl(p)
            for r in p.runs:
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                set_rtl(p)
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(10)
    doc.add_paragraph()


def build_user_guide(doc):
    add_heading(doc, "חלק א׳ — מדריך למשתמש", level=1)

    add_heading(doc, "1. מבוא", level=2)
    add_para(doc,
        "מערכת TripLink היא מערכת מבוססת אינטרנט לתכנון ואופטימיזציה של טיולים יומיים "
        "בתחבורה ציבורית. המערכת מאפשרת למשתמש ליצור טיול מותאם אישית, לבחור העדפות "
        "(אזור, קטגוריות, מאפיינים, זמנים), לחשב מסלול אופטימלי ולצפות בתוצאה — מפה, "
        "ציר זמן ופרטי נסיעה.")
    add_para(doc, "קהל יעד: מטיילים המתכננים יום טיול בישראל באמצעות אוטובוסים ותחבורה ציבורית.")

    add_heading(doc, "2. דרישות מערכת", level=2)
    add_table(doc, ["רכיב", "דרישה"],
        [
            ("דפדפן", "Chrome, Edge או Firefox (גרסה עדכנית)"),
            ("חיבור רשת", "גישה ל-localhost בפיתוח, או לשרת הפרודקשן"),
            ("שרת API", ".NET — http://localhost:3000"),
            ("אפליקציית Web", "Angular — http://localhost:4200"),
            ("מפת Google (אופציונלי)", "מפתח API ב-environment.ts להצגת מפה במסך תוצאות"),
        ])

    add_heading(doc, "3. כניסה למערכת והרשמה", level=2)

    add_heading(doc, "3.1 הרשמת משתמש חדש", level=3)
    add_numbered(doc, "פתחי את הדפדפן וגשי לכתובת: http://localhost:4200")
    add_numbered(doc, "במסך ההתחברות, לחצי על «הירשם כאן».")
    add_numbered(doc, "מלאי: שם מלא, אימייל, טלפון וסיסמה.")
    add_numbered(doc, "לחצי «הירשם» — לאחר הצלחה תועברי למסך ההתחברות.")

    add_heading(doc, "3.2 התחברות", level=3)
    add_numbered(doc, "הזיני אימייל וסיסמה.")
    add_numbered(doc, "לחצי «התחבר».")
    add_numbered(doc, "בהצלחה — מעבר אוטומטי למסך «הטיולים שלי».")
    add_para(doc, "שגיאות אפשריות: «שגיאה בהתחברות» — אימייל או סיסמה שגויים.", bold=False)

    add_heading(doc, "4. מסך הטיולים שלי (/my-trips)", level=2)
    add_para(doc, "מסך הבית לאחר התחברות. מציג את כל הטיולים של המשתמש המחובר.")
    add_table(doc, ["פעולה", "איך לבצע", "תוצאה"],
        [
            ("צפייה ברשימת טיולים", "המסך נטען אוטומטית", "טבלה: מזהה, שם, תאריך, כתובת התחלה"),
            ("תכנון טיול חדש", "לחצי «+ תכנן טיול חדש»", "מעבר לאשף התכנון (/plan)"),
            ("צפייה במסלול", "לחצי «צפה במסלול» בשורת טיול", "מסך תוצאות (/trips/{id}/result)"),
            ("מחיקת טיול", "לחצי «מחק» ואשרי", "הטיול נמחק מהרשימה"),
        ])

    add_heading(doc, "5. אשף תכנון טיול חדש (/plan) — 3 שלבים", level=2)

    add_heading(doc, "שלב 1 — פרטים", level=3)
    add_table(doc, ["שדה", "חובה", "הסבר"],
        [
            ("שם הטיול", "כן", "לדוגמה: טיול דרום קלאסי"),
            ("אזור טיול", "כן", "דרום / צפון / מרכז"),
            ("תאריך הטיול", "לא", "תאריך מתוכנן"),
            ("כתובת התחלה", "לא", "לדוגמה: תחנה מרכזית באר שבע"),
        ])
    add_para(doc, "לחצי «הבא» — לא ניתן להמשיך בלי שם ואזור.")

    add_heading(doc, "שלב 2 — העדפות", level=3)
    add_table(doc, ["שדה", "חובה", "הסבר"],
        [
            ("רמת קושי", "לא", "קל / בינוני / קשה"),
            ("קטגוריות", "לא", "סימון מרובה: טבע, היסטוריה, חוף וכו'"),
            ("מאפיינים נדרשים", "לא", "חניה חינם, נגישות, כניסה חינם וכו'"),
            ("מינימום / מקסימום יעדים", "לא", "טווח יעדים במסלול"),
        ])
    add_para(doc, "לחצי «הבא».")

    add_heading(doc, "שלב 3 — זמנים", level=3)
    add_table(doc, ["שדה", "ברירת מחדל", "הסבר"],
        [
            ("שעת יציאה", "08:00", "תחילת יום הטיול"),
            ("שעת חזרה", "18:00", "סיום יום הטיול"),
            ("זמן נסיעה מקסימלי (דקות)", "480", "תקציב נסיעה לאופטימייזר"),
            ("זמן חזרה (דקות)", "60", "מרווח לפני הוספת יעד (בשרת: שעות)"),
            ("יעילות תחבורה מינימלית", "0.5", "סף 0–1 לסינון נסיעות"),
        ])
    add_para(doc, "לחצי «צור טיול והמשך לאופטימיזציה» — נוצר טיול ב-DB ומעבר למסך האופטימיזציה.")

    add_heading(doc, "6. מסך אופטימיזציה (/plan/optimize/:tripId)", level=2)
    add_numbered(doc, "מוצגים סיכום הטיול ושעות.")
    add_numbered(doc, "לחצי «חשב מסלול אופטימלי».")
    add_numbered(doc, "המערכת מפעילה את אלגוריתם האופטימיזציה בשרת.")
    add_numbered(doc, "בהצלחה — שמירת המסלול ומעבר למסך התוצאות.")
    add_para(doc, "שגיאה אפשרית: «שגיאה בחישוב המסלול» — אין יעדים מתאימים או בעיית שרת.")

    add_heading(doc, "7. מסך תוצאות מסלול (/trips/:id/result)", level=2)
    add_para(doc, "מציג את תוצאת האלגוריתם:")
    add_bullet(doc, "כותרת, נקודת התחלה, סטטיסטיקות (יעדים, ציון, יעילות)")
    add_bullet(doc, "מפת Google עם נקודות וקו מסלול (דורש מפתח API)")
    add_bullet(doc, "ציר זמן — כרטיס לכל יעד: תמונה, שעות, פרטי אוטובוס")
    add_bullet(doc, "סיכום מילולי של המסלול")
    add_bullet(doc, "כפתורים: חזרה לטיולים שלי / תכנון טיול חדש")

    add_heading(doc, "8. מסך יעדים (/destinations)", level=2)
    add_para(doc, "קטלוג כל היעדים במערכת — שם, אזור, רמת קושי, סוג מטייל, זמן ביקור. לצפייה בלבד.")

    add_heading(doc, "9. זרימת עבודה מלאה (תמצית)", level=2)
    add_para(doc,
        "התחברות → הטיולים שלי → תכנן טיול חדש → שלב 1 (פרטים) → שלב 2 (העדפות) → "
        "שלב 3 (זמנים) → חשב מסלול → צפייה בתוצאה → (אופציונלי) חזרה לטיולים שלי")

    add_heading(doc, "10. פתרון בעיות נפוצות", level=2)
    add_table(doc, ["בעיה", "סיבה אפשרית", "פתרון"],
        [
            ("שגיאה ביצירת טיול", "שרת API לא רץ", "הפעילי: dotnet run בתיקיית API"),
            ("מסך ריק בתוצאות", "לא הורצה אופטימיזציה", "חזרי לאשף והריצי אופטימיזציה"),
            ("מפה לא מוצגת", "אין מפתח Google Maps", "הגדירי googleMapsApiKey ב-environment.ts"),
            ("אין יעדים במסלול", "אין יעדים באזור/העדפות", "שני אזור או הרחיבי קטגוריות"),
            ("שגיאת התחברות", "משתמש לא רשום", "הירשמי או בדקי אימייל/סיסמה"),
        ])


def build_testing(doc):
    add_heading(doc, "חלק ב׳ — בדיקות והערכה", level=1)

    add_heading(doc, "1. מטרת הבדיקות", level=2)
    add_para(doc,
        "מטרת פרק הבדיקות היא לוודא שהמערכת עומדת בדרישות הפונקציונליות, "
        "שזרימת המשתמש שלמה, ושתוצאות האלגוריתם הגיוניות. "
        "הבדיקות מחולקות לבדיקות יחידה (שרת), בדיקות אינטגרציה (API) "
        "ובדיקות מערכת (E2E — ממשק משתמש).")

    add_heading(doc, "2. סביבת בדיקה", level=2)
    add_table(doc, ["פרמטר", "ערך"],
        [
            ("מסד נתונים", "SQL Server — TripLink (עם SeedData)"),
            ("API", "http://localhost:3000"),
            ("Frontend", "http://localhost:4200"),
            ("דפדפן", "Chrome / Edge"),
            ("משתמש בדיקה", "נרשם חדש או משתמש מ-SeedData"),
        ])

    add_heading(doc, "3. בדיקות פונקציונליות — ממשק משתמש", level=2)
    add_para(doc, "טבלת מקרי בדיקה עיקריים:")

    test_cases = [
        ("TC-01", "הרשמה", "מילוי טופס הרשמה תקין", "משתמש נוצר, מעבר להתחברות", "עבר / נכשל"),
        ("TC-02", "התחברות", "אימייל וסיסמה נכונים", "מעבר ל-/my-trips", "עבר / נכשל"),
        ("TC-03", "התחברות שלילית", "סיסמה שגויה", "הודעת שגיאה, נשאר ב-login", "עבר / נכשל"),
        ("TC-04", "הגנת נתיב", "גישה ל-/plan ללא התחברות", "הפניה ל-login", "עבר / נכשל"),
        ("TC-05", "יצירת טיול", "אשף 3 שלבים — דרום, 2 קטגוריות", "טיול נוצר, מעבר לאופטימיזציה", "עבר / נכשל"),
        ("TC-06", "ולידציה שלב 1", "ללא שם או אזור", "כפתור «הבא» מושבת", "עבר / נכשל"),
        ("TC-07", "אופטימיזציה", "לחיצה «חשב מסלול»", "מסלול עם יעדים, מעבר לתוצאות", "עבר / נכשל"),
        ("TC-08", "תוצאות", "צפייה במסלול אחרי אופטימיזציה", "יעדים, שעות, narrative", "עבר / נכשל"),
        ("TC-09", "רשימת יעדים", "גישה ל-/destinations", "טבלה עם 10 יעדים", "עבר / נכשל"),
        ("TC-10", "מחיקת טיול", "מחק + אישור", "טיול נעלם מהרשימה", "עבר / נכשל"),
        ("TC-11", "צפייה במסלול קיים", "«צפה במסלול» מ-my-trips", "טעינת itinerary מ-API", "עבר / נכשל"),
        ("TC-12", "טיול ללא מסלול", "result לטיול שלא אופטם", "הודעת שגיאה מתאימה", "עבר / נכשל"),
    ]
    add_table(doc, ["מזהה", "נושא", "קלט", "פלט צפוי", "תוצאה"],
              test_cases)

    add_heading(doc, "4. בדיקות API (אינטגרציה)", level=2)
    api_tests = [
        ("API-01", "POST /api/auth/register", "פרטים תקינים", "200 + token"),
        ("API-02", "POST /api/auth/login", "אימייל/סיסמה", "200 + JWT"),
        ("API-03", "POST /api/trips", "CreateTripDto מלא", "201 + TripId"),
        ("API-04", "GET /api/trips/user/{id}", "userId קיים", "רשימת טיולים"),
        ("API-05", "POST /api/trips/optimize", "OptimizeRequestDto", "OptimizeResultDto"),
        ("API-06", "GET /api/trips/{id}/itinerary", "טיול עם מסלול", "TripItineraryDto"),
        ("API-07", "GET /api/destinations", "—", "רשימת יעדים"),
        ("API-08", "DELETE /api/trips/{id}", "tripId קיים", "204"),
    ]
    add_table(doc, ["מזהה", "Endpoint", "קלט", "פלט צפוי"], api_tests)

    add_heading(doc, "5. בדיקות אלגוריתם האופטימיזציה", level=2)
    add_table(doc, ["מזהה", "תרחיש", "קלט", "קריטריון הצלחה"],
        [
            ("OPT-01", "חלון זמן", "08:00–18:00", "כל היעדים בתוך חלון הזמן"),
            ("OPT-02", "אזור דרום", "region=דרום", "רק יעדים מאזור דרום"),
            ("OPT-03", "מקסימום יעדים", "maxNumDes=3", "לא יותר מ-3 יעדים"),
            ("OPT-04", "קטגוריות", "קטגוריה היסטוריה", "יעדים מקטגוריה זו בלבד"),
            ("OPT-05", "יעילות תחבורה", "minTransitEfficiency=0.7", "יעילות ממוצעת ≥ 0.7"),
            ("OPT-06", "מסלול ריק", "אזור ללא יעדים", "הודעת שגיאה / 0 יעדים"),
        ])

    add_heading(doc, "6. בדיקות לא-פונקציונליות", level=2)
    add_table(doc, ["קטגוריה", "בדיקה", "קריטריון"],
        [
            ("ביצועים", "אופטימיזציה ל-10 יעדים", "תגובה < 10 שניות"),
            ("שימושיות", "אשף 3 שלבים", "משתמש חדש מסיים ללא עזרה"),
            ("אבטחה", "גישה ללא JWT", "401 / הפניה ל-login"),
            ("תאימות", "RTL עברית", "טקסט וטבלאות מיושרים ימינה"),
            ("שחזור", "רענון בדף תוצאות", "טעינה מ-API אם אין cache"),
        ])

    add_heading(doc, "7. הערכה ומדדי הצלחה", level=2)
    add_para(doc, "הערכת המערכת מתבססת על מדדים כמותיים ואיכותיים:")

    add_heading(doc, "7.1 מדדים כמותיים", level=3)
    add_table(doc, ["מדד", "הגדרה", "יעד"],
        [
            ("שיעור הצלחת בדיקות", "TC שעברו / סה\"כ TC", "≥ 90%"),
            ("זמן תגובת API", "ממוצע בקשות עיקריות", "< 3 שניות"),
            ("יעדי אופטימיזציה", "ממוצע יעדים במסלול", "2–4 יעדים"),
            ("יעילות תחבורה", "TransitEfficiency ממוצע", "≥ 50%"),
        ])

    add_heading(doc, "7.2 מדדים איכותיים", level=3)
    add_bullet(doc, "האם המסלול הגיוני מבחינת שעות וסדר יעדים?")
    add_bullet(doc, "האם הממשק ברור למשתמש ללא הדרכה?")
    add_bullet(doc, "האם פרטי האוטובוס וההליכה מסייעים בתכנון?")
    add_bullet(doc, "האם השגיאות מוצגות בעברית ומובנות?")

    add_heading(doc, "7.3 שאלון משתמש (דוגמה)", level=3)
    add_table(doc, ["שאלה", "סולם 1–5"],
        [
            ("האם המערכת קלה לשימוש?", "1=קשה מאוד, 5=קל מאוד"),
            ("האם המסלול שקיבלתי היה שימושי?", "1=לא, 5=מאוד"),
            ("האם היית משתמשת במערכת שוב?", "1=לא, 5=בהחלט"),
            ("האם המידע על תחבורה ציבורית היה מספיק?", "1=לא, 5=מאוד"),
        ])

    add_heading(doc, "8. סיכום בדיקות — תבנית למילוי", level=2)
    add_table(doc, ["קטגוריה", "נבדקו", "עברו", "נכשלו", "אחוז הצלחה"],
        [
            ("בדיקות UI (TC)", "12", "", "", ""),
            ("בדיקות API", "8", "", "", ""),
            ("בדיקות אופטימיזציה", "6", "", "", ""),
            ("בדיקות לא-פונקציונליות", "5", "", "", ""),
            ("סה\"כ", "31", "", "", ""),
        ])
    add_para(doc, "הערות בודק: _______________________________________________")
    add_para(doc, "תאריך בדיקה: _____________  חתימה: _____________")


def build():
    doc = Document()
    add_heading(doc, "TripLink — מדריך למשתמש, בדיקות והערכה", level=1)
    add_para(doc, "מסמך לספר פרויקט — גרסה 1.0")
    doc.add_paragraph()

    build_user_guide(doc)
    doc.add_page_break()
    build_testing(doc)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
